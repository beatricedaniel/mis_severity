# text_extraction.py
import docx
import json
from typing import List

# Load settings
with open("settings.json", "r") as f:
    settings = json.load(f)

def extract_text_from_paragraphs(doc: docx.Document) -> List[str]:
    """Extracts text from paragraphs in a DOCX document."""
    return [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]

def extract_text_from_tables(doc: docx.Document) -> List[str]:
    """Extracts text from tables in a DOCX document."""
    table_texts = []
    for table in doc.tables:
        for row in table.rows:
            row_text = " ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                table_texts.append(row_text)
    return table_texts

def extract_text_from_docx(docx_path: str) -> List[str]:
    """Extracts text from a DOCX document, including paragraphs and tables."""
    doc = docx.Document(docx_path)
    paragraphs = extract_text_from_paragraphs(doc)
    tables = extract_text_from_tables(doc)
    return paragraphs + tables

def extract_text(file_path: str, file_type: str = "docx") -> List[str]:
    """Extracts text based on file type (currently supports DOCX)."""
    if file_type.lower() == "docx":
        return extract_text_from_docx(file_path)
    else:
        raise ValueError("Unsupported file type. Only 'docx' is supported in this configuration.")
